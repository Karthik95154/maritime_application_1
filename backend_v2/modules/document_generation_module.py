import os
import json
from google import genai

from dotenv import load_dotenv

from docx import Document
from docx.shared import Inches
from docx.shared import Pt


class DocumentGenerationModule:

    def __init__(

        self,

        gemini_model_name="gemini-1.5-flash",

        output_folder="outputs/final_reports"
    ):

        load_dotenv()

        api_key = os.getenv(
            "GEMINI_API_KEY"
        )

        self.client = genai.Client(
            api_key=api_key
        )

        self.gemini_model_name = gemini_model_name

        self.output_folder = output_folder

        os.makedirs(
            self.output_folder,
            exist_ok=True
        )

    # =====================================================
    # JSON HELPERS
    # =====================================================

    def load_json(self, path):

        with open(path, "r") as f:
            return json.load(f)

    # =====================================================
    # LLM CALL
    # =====================================================

    def generate_text(self, prompt):
        import time
        max_retries = 5
        for i in range(max_retries):
            try:
                response = self.client.models.generate_content(
                    model=self.gemini_model_name,
                    contents=prompt
                )
                return response.text
            except Exception as e:
                err_str = str(e)
                print(f"[ERROR] Gemini generation failed: {err_str}")
                if "429" in err_str or "Quota" in err_str or "exceeded" in err_str.lower():
                    sleep_time = 30 * (i + 1)
                    print(f"[INFO] Rate limit hit. Sleeping for {sleep_time} seconds before retry {i+1}/{max_retries}...")
                    time.sleep(sleep_time)
                else:
                    return "Generation failed."
        return "Generation failed after retries."

    # =====================================================
    # EXECUTIVE SUMMARY
    # =====================================================

    def generate_executive_summary(

        self,

        repair_summary
    ):

        prompt = f"""
        Generate a professional maritime vessel inspection
        executive summary.

        Total defects: {repair_summary['total_defects']}

        Total estimated repair cost:
        {repair_summary['total_estimated_cost']}
        {repair_summary['currency']}

        Severity distribution:
        {repair_summary['severity_distribution']}

        Keep the explanation professional,
        concise and technical.
        """

        return self.generate_text(prompt)

    # =====================================================
    # DEFECT EXPLANATION
    # =====================================================

    def generate_defect_explanation(

        self,

        repair_data
    ):

        prompt = f"""
        Explain the following maritime defect
        professionally for a vessel inspection report.

        Defect Name:
        {repair_data['defect_name']}

        Severity:
        {repair_data['severity']}

        Description:
        {repair_data['description']}

        Overlapping Parts:
        {repair_data['defect_metadata']['overlapping_parts']}

        Explain:

        1. What the defect means.
        2. Why it is important.
        3. Risks if not repaired.
        4. Structural implications.

        Keep the explanation clear,
        professional and easy to understand.
        """

        return self.generate_text(prompt)

    # =====================================================
    # REPAIR EXPLANATION
    # =====================================================

    def generate_repair_explanation(

        self,

        repair_data
    ):

        prompt = f"""
        Explain the following maritime repair process
        professionally for a vessel inspection report.

        Repair Process:
        {repair_data['repair_process']}

        Required Items:
        {repair_data['repair_estimation']['required_items']}

        Explain:

        1. Why each repair step is necessary.
        2. What happens during the repair.
        3. Why the required items are used.
        4. Safety and structural considerations.

        Keep the explanation professional.
        """

        return self.generate_text(prompt)

    # =====================================================
    # CREATE REPORT
    # =====================================================

    def create_report(

        self,

        repair_estimation_json_path
    ):

        repair_outputs = self.load_json(
            repair_estimation_json_path
        )

        repair_summary = (
            repair_outputs[
                "repair_summary"
            ]
        )

        defect_repairs = (
            repair_outputs[
                "defect_repairs"
            ]
        )

        # =================================================
        # DOCX
        # =================================================

        document = Document()

        # =================================================
        # TITLE
        # =================================================

        title = document.add_heading(
            "Vessel Inspection Report",
            level=0
        )

        title.runs[0].font.size = Pt(24)

        # =================================================
        # TABLE OF CONTENTS
        # =================================================

        document.add_heading(
            "Table of Contents",
            level=1
        )

        toc_items = [

            "1. Executive Summary",
            "2. Vessel Repair Summary",
            "3. Defect Analysis"
        ]

        for idx, defect_id in enumerate(
            defect_repairs.keys(),
            start=1
        ):

            toc_items.append(
                f"3.{idx} {defect_id}"
            )

        for item in toc_items:
            document.add_paragraph(item)

        # =================================================
        # EXECUTIVE SUMMARY
        # =================================================

        document.add_page_break()

        document.add_heading(
            "1. Executive Summary",
            level=1
        )

        executive_summary = (
            self.generate_executive_summary(
                repair_summary
            )
        )

        document.add_paragraph(
            executive_summary
        )

        # =================================================
        # REPAIR SUMMARY
        # =================================================

        document.add_heading(
            "2. Vessel Repair Summary",
            level=1
        )

        summary_table = document.add_table(
            rows=1,
            cols=2
        )

        summary_table.style = "Table Grid"

        hdr_cells = summary_table.rows[0].cells

        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"

        summary_items = {

            "Total Defects":
                repair_summary[
                    "total_defects"
                ],

            "Total Estimated Cost":
                f"{repair_summary['total_estimated_cost']} "
                f"{repair_summary['currency']}",

            "Material Cost":
                f"{repair_summary['total_material_cost']} "
                f"{repair_summary['currency']}",

            "Labor Cost":
                f"{repair_summary['total_labor_cost']} "
                f"{repair_summary['currency']}",

            "Equipment Cost":
                f"{repair_summary['total_equipment_cost']} "
                f"{repair_summary['currency']}"
        }

        for key, value in summary_items.items():

            row_cells = summary_table.add_row().cells

            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

        # =================================================
        # DEFECT ANALYSIS
        # =================================================

        document.add_page_break()

        document.add_heading(
            "3. Defect Analysis",
            level=1
        )

        for idx, (
            persistent_id,
            repair_data
        ) in enumerate(
            defect_repairs.items(),
            start=1
        ):

            # =============================================
            # SECTION TITLE
            # =============================================

            document.add_heading(
                f"3.{idx} {persistent_id}",
                level=2
            )

            # =============================================
            # DEFECT OVERVIEW
            # =============================================

            document.add_heading(
                "Defect Overview",
                level=3
            )

            defect_explanation = (
                self.generate_defect_explanation(
                    repair_data
                )
            )

            document.add_paragraph(
                defect_explanation
            )

            # =============================================
            # VISUAL EVIDENCE
            # =============================================

            document.add_heading(
                "Visual Evidence",
                level=3
            )

            best_frame_path = (
                repair_data[
                    "defect_metadata"
                ][
                    "best_frame_path"
                ]
            )

            if os.path.exists(best_frame_path):

                document.add_picture(
                    best_frame_path,
                    width=Inches(5.5)
                )

            # =============================================
            # REPAIR EXPLANATION
            # =============================================

            document.add_heading(
                "Repair Procedure",
                level=3
            )

            repair_explanation = (
                self.generate_repair_explanation(
                    repair_data
                )
            )

            document.add_paragraph(
                repair_explanation
            )

            # =============================================
            # COST TABLE
            # =============================================

            document.add_heading(
                "Repair Cost Breakdown",
                level=3
            )

            cost_table = document.add_table(
                rows=1,
                cols=5
            )

            cost_table.style = "Table Grid"

            hdr = cost_table.rows[0].cells

            hdr[0].text = "Item"
            hdr[1].text = "Quantity"
            hdr[2].text = "Unit Cost"
            hdr[3].text = "Total Cost"
            hdr[4].text = "Currency"

            required_items = (
                repair_data[
                    "repair_estimation"
                ][
                    "required_items"
                ]
            )

            for item in required_items:

                row = cost_table.add_row().cells

                row[0].text = str(
                    item["item_name"]
                )

                row[1].text = str(
                    item["required_quantity"]
                )

                row[2].text = str(
                    item["unit_cost"]
                )

                row[3].text = str(
                    item["total_cost"]
                )

                row[4].text = str(
                    item["currency"]
                )

            document.add_paragraph(
                f"Estimated Repair Cost: "
                f"{repair_data['repair_estimation']['estimated_total_cost']} "
                f"{repair_data['repair_estimation']['currency']}"
            )

        # =================================================
        # SAVE DOCX
        # =================================================

        output_docx_path = os.path.join(

            self.output_folder,

            "vessel_inspection_report.docx"
        )

        document.save(
            output_docx_path
        )

        print(
            f"[INFO] Report saved: "
            f"{output_docx_path}"
        )

        return output_docx_path

    # =====================================================
    # CREATE BATCH REPORT
    # =====================================================

    def create_batch_report(self, batch_id: str, repair_json_paths: list[str], vessel_name: str):
        aggregated_total_defects = 0
        aggregated_total_cost = 0.0
        aggregated_material_cost = 0.0
        aggregated_labor_cost = 0.0
        aggregated_equipment_cost = 0.0
        currency = "INR"
        all_defect_repairs = {}
        
        for path in repair_json_paths:
            if not os.path.exists(path): continue
            repair_outputs = self.load_json(path)
            repair_summary = repair_outputs.get("repair_summary", {})
            
            aggregated_total_defects += int(repair_summary.get("total_defects", 0))
            aggregated_total_cost += float(repair_summary.get("total_estimated_cost", 0))
            aggregated_material_cost += float(repair_summary.get("total_material_cost", 0))
            aggregated_labor_cost += float(repair_summary.get("total_labor_cost", 0))
            aggregated_equipment_cost += float(repair_summary.get("total_equipment_cost", 0))
            currency = repair_summary.get("currency", "INR")
            
            defect_repairs = repair_outputs.get("defect_repairs", {})
            for d_id, d_data in defect_repairs.items():
                all_defect_repairs[d_id] = d_data

        combined_summary = {
            "total_defects": aggregated_total_defects,
            "total_estimated_cost": aggregated_total_cost,
            "total_material_cost": aggregated_material_cost,
            "total_labor_cost": aggregated_labor_cost,
            "total_equipment_cost": aggregated_equipment_cost,
            "currency": currency,
            "severity_distribution": "Mixed severity across batch"
        }

        document = Document()
        title = document.add_heading(f"Batch Inspection Report: {vessel_name}", level=0)
        title.runs[0].font.size = Pt(24)

        document.add_heading("Table of Contents", level=1)
        toc_items = ["1. Executive Summary", "2. Batch Repair Summary", "3. Defect Analysis"]
        for idx, defect_id in enumerate(all_defect_repairs.keys(), start=1):
            toc_items.append(f"3.{idx} {defect_id}")
        for item in toc_items:
            document.add_paragraph(item)

        document.add_page_break()
        document.add_heading("1. Executive Summary", level=1)
        executive_summary = self.generate_executive_summary(combined_summary)
        document.add_paragraph(executive_summary)

        document.add_heading("2. Batch Repair Summary", level=1)
        summary_table = document.add_table(rows=1, cols=2)
        summary_table.style = "Table Grid"
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"
        
        summary_items = {
            "Total Defects": combined_summary["total_defects"],
            "Total Estimated Cost": f"{combined_summary['total_estimated_cost']} {currency}",
            "Material Cost": f"{combined_summary['total_material_cost']} {currency}",
            "Labor Cost": f"{combined_summary['total_labor_cost']} {currency}",
            "Equipment Cost": f"{combined_summary['total_equipment_cost']} {currency}"
        }
        for key, value in summary_items.items():
            row_cells = summary_table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

        document.add_page_break()
        document.add_heading("3. Defect Analysis", level=1)

        for idx, (persistent_id, repair_data) in enumerate(all_defect_repairs.items(), start=1):
            document.add_heading(f"3.{idx} {persistent_id}", level=2)
            
            document.add_heading("Defect Overview", level=3)
            defect_explanation = self.generate_defect_explanation(repair_data)
            document.add_paragraph(defect_explanation)

            document.add_heading("Visual Evidence", level=3)
            best_frame_path = repair_data.get("defect_metadata", {}).get("best_frame_path", "")
            if os.path.exists(best_frame_path):
                document.add_picture(best_frame_path, width=Inches(5.5))

            document.add_heading("Repair Procedure", level=3)
            repair_explanation = self.generate_repair_explanation(repair_data)
            document.add_paragraph(repair_explanation)

            document.add_heading("Repair Cost Breakdown", level=3)
            cost_table = document.add_table(rows=1, cols=5)
            cost_table.style = "Table Grid"
            hdr = cost_table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Item", "Quantity", "Unit Cost", "Total Cost", "Currency"
            
            required_items = repair_data.get("repair_estimation", {}).get("required_items", [])
            for item in required_items:
                row = cost_table.add_row().cells
                row[0].text = str(item.get("item_name", ""))
                row[1].text = str(item.get("required_quantity", ""))
                row[2].text = str(item.get("unit_cost", ""))
                row[3].text = str(item.get("total_cost", ""))
                row[4].text = str(item.get("currency", ""))

            est_cost = repair_data.get("repair_estimation", {}).get("estimated_total_cost", 0)
            est_curr = repair_data.get("repair_estimation", {}).get("currency", currency)
            document.add_paragraph(f"Estimated Repair Cost: {est_cost} {est_curr}")

        output_dir = os.path.join("outputs", "batches", batch_id)
        os.makedirs(output_dir, exist_ok=True)
        output_docx_path = os.path.join(output_dir, "combined_vessel_inspection_report.docx")
        
        document.save(output_docx_path)
        print(f"[INFO] Batch Report saved: {output_docx_path}")
        return output_docx_path

# =========================================================
# TESTING
# =========================================================
"""
if __name__ == "__main__":

    document_generator = (
        DocumentGenerationModule(
            gemini_model_name="gemini-2.5-flash",
            output_folder="frame_extraction_testing_outputs/deformation_3/document_generation_output/"
        )
    )

    document_generator.create_report(

        repair_estimation_json_path=
            "frame_extraction_testing_outputs/deformation_3/repair_estimation_output/repair_estimation_outputs.json"
    )
"""